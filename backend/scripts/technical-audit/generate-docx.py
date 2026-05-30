#!/usr/bin/env python3
"""Generate technical SEO audit DOCX from unified JSON report."""

from __future__ import annotations

import argparse
import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor

STATUS_LABELS = {
    "ok": "Все в порядке.",
    "warning": "Есть замечания.",
    "critical": "Есть проблема.",
}

STATUS_INTRO = {
    "ok": "Все в порядке. В пунктах технического аудита, обозначенных этим знаком, ошибок не обнаружено.",
    "warning": (
        "Есть замечания. В пунктах технического аудита, обозначенных этим знаком, "
        "присутствуют замечания, которые не имеют критичного для продвижения сайта значения. "
        "Устранение этих ошибок не является приоритетным."
    ),
    "critical": (
        "Есть проблема. В пунктах технического аудита, обозначенных этим знаком, "
        "обнаружены грубые технические нарушения. Устранение этих ошибок является первостепенной задачей."
    ),
}

SEO_FACTOR_HINTS: dict[str, str] = {
    "main_mirror": (
        "Главное зеркало сайта должно соответствовать форме записи домена в браузере. "
        "Расхождение может привести к выпадению страниц из индекса."
    ),
    "robots_txt_content": (
        "Файл robots.txt задаёт параметры индексирования. "
        "Он должен закрывать мусорные страницы и не блокировать продвигаемые разделы."
    ),
    "duplicate_title": (
        "Тег Title — самый весомый текстовый элемент страницы. "
        "Дублирование размывает релевантность и снижает позиции."
    ),
    "broken_links": (
        "Битые ссылки негативно влияют на поведенческие факторы и доверие поисковых систем к сайту."
    ),
}


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str, bold: bool = False) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)


def add_status_paragraph(doc: Document, status: str) -> None:
    color_map = {
        "ok": RGBColor(0, 128, 0),
        "warning": RGBColor(204, 102, 0),
        "critical": RGBColor(192, 0, 0),
    }
    p = doc.add_paragraph()
    run = p.add_run(STATUS_LABELS.get(status, status))
    run.bold = True
    run.font.color.rgb = color_map.get(status, RGBColor(0, 0, 0))


def seo_factor_text(check: dict) -> str:
    if check.get("seo_factor"):
        return check["seo_factor"]
    check_id = check.get("id", "")
    if check_id in SEO_FACTOR_HINTS:
        return SEO_FACTOR_HINTS[check_id]
    return (
        check.get("finding", "")
        or "Фактор влияет на индексацию, релевантность или доверие поисковых систем к сайту."
    )


def build_document(data: dict) -> Document:
    doc = Document()
    site_url = data.get("site_url", "")
    site_name = data.get("site_name") or site_url
    audit_date = data.get("audit_date", "")

    add_heading(doc, "Технический аудит сайта", level=0)
    if site_name:
        add_paragraph(doc, f"Сайт: {site_url}")
        add_paragraph(doc, f"Компания: {site_name}")
    if audit_date:
        add_paragraph(doc, f"Дата аудита: {audit_date}")

    add_heading(doc, "Введение", level=1)
    intro = [
        (
            "Цель технического аудита сайта — выявить критические ошибки, "
            "мешающие корректному индексированию и продвижению сайта в поисковых системах, "
            "и способы их устранения."
        ),
        (
            "Выполнение этих рекомендаций имеет высокий приоритет, так как их несоблюдение "
            "может создавать очень серьёзные препятствия для работ по продвижению сайта."
        ),
        (
            "Данный отчёт содержит перечень и описание обнаруженных ошибок и не является "
            "готовым техническим заданием для конечных исполнителей."
        ),
        (
            "По результату проведения технического аудита техническим специалистам "
            "будут сформированы задачи на устранение обнаруженных технических ошибок."
        ),
    ]
    for paragraph in intro:
        add_paragraph(doc, paragraph)

    add_heading(doc, "Обозначения статусов", level=1)
    for key in ("ok", "warning", "critical"):
        add_paragraph(doc, STATUS_INTRO[key], bold=True)

    checks = data.get("checks") or []
    totals = data.get("totals") or {}

    add_heading(doc, "Оглавление", level=1)
    for i, check in enumerate(checks, 1):
        title = check.get("title") or check.get("id") or f"Пункт {i}"
        doc.add_paragraph(f"{i}. {title}", style="List Number")

    add_heading(doc, "Сводка", level=1)
    add_paragraph(
        doc,
        f"Критичных: {totals.get('critical', 0)} | "
        f"Замечаний: {totals.get('warning', 0)} | "
        f"Без проблем: {totals.get('ok', 0)}",
        bold=True,
    )

    priorities = data.get("top_priorities") or []
    if priorities:
        add_heading(doc, "Приоритет исправлений", level=1)
        for item in priorities:
            doc.add_paragraph(str(item), style="List Bullet")

    for check in checks:
        title = check.get("title") or check.get("id") or "Проверка"
        status = check.get("status", "warning")
        finding = check.get("finding") or "Данные не предоставлены."
        evidence = check.get("evidence") or []

        add_heading(doc, title, level=1)
        add_status_paragraph(doc, status)
        add_paragraph(doc, finding)

        if evidence:
            add_paragraph(doc, "Доказательства / примеры:", bold=True)
            for item in evidence[:30]:
                doc.add_paragraph(str(item), style="List Bullet")
            if len(evidence) > 30:
                add_paragraph(doc, f"... и ещё {len(evidence) - 30} записей (см. JSON).")

        add_paragraph(doc, "Об этом SEO факторе:", bold=True)
        add_paragraph(doc, seo_factor_text(check))

    metrics = data.get("metrics") or {}
    if metrics:
        add_heading(doc, "Метрики on-page", level=1)
        for key, value in metrics.items():
            add_paragraph(doc, f"{key}: {value}")

    return doc


def main() -> int:
    parser = argparse.ArgumentParser(description="Generate SEO audit DOCX from JSON")
    parser.add_argument("--input", required=True, help="Path to unified audit JSON")
    parser.add_argument("--output", help="Output DOCX path (default: same name as input)")
    args = parser.parse_args()

    input_path = Path(args.input)
    if not input_path.exists():
        print(f"Input file not found: {input_path}", file=sys.stderr)
        return 1

    output_path = Path(args.output) if args.output else input_path.with_suffix(".docx")

    with input_path.open(encoding="utf-8") as f:
        data = json.load(f)

    doc = build_document(data)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(str(output_path.resolve()))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
